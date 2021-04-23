import numpy as np
import sys 
import matplotlib.pyplot as plot
import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LinearRegression
from sklearn.metrics import r2_score
from sklearn.linear_model import Ridge
from sklearn import linear_model
from sklearn.linear_model import ElasticNet
from sklearn.datasets import make_regression
import statsmodels.api as sm
from statsmodels.sandbox.regression.predstd import wls_prediction_std
from sklearn.metrics import mean_absolute_error
from sklearn.metrics import mean_squared_error
from matplotlib.pyplot import figure

a1 = sys.argv[1]  # File CSV
a2 = float(sys.argv[2]) # Train/Test Ration it should be float number less than 1
a3 = sys.argv[3] #Normalize True/False

dataset = pd.read_csv(a1)

X=dataset.iloc[:,1:]
Y=dataset.iloc[:,0]

xTrain, xTest, yTrain, yTest = train_test_split(X, Y, test_size = a2, random_state = 0)

linearRegressor = LinearRegression(normalize=a3,fit_intercept=True)
linearRegressor.fit(xTrain, yTrain)

yPredictionTest = linearRegressor.predict(xTest)
yPredictionTrain = linearRegressor.predict(xTrain)
#print(type(yTest))
#print(type(yPredictionTest))

figure(num=None, figsize=(5, 4), dpi=80, facecolor='w', edgecolor='k')
plot.scatter(yTrain,yPredictionTrain, color = 'red')
#plot.plot(yTrain, linearRegressor.predict(xTrain), color = 'blue')
plot.title('Train Set Plot')
plot.xlabel('Observed Activity')
plot.ylabel('Predicted Activity')
plot.savefig('TrainSet.png')
plot.show()

figure(num=None, figsize=(5, 4), dpi=80, facecolor='w', edgecolor='k')
plot.scatter(yTest,yPredictionTest, color = 'red')
#plot.plot(yTest, linearRegressor.predict(xTest), color = 'blue')
plot.title('Test Set Plot')
plot.xlabel('Observed Activity')
plot.ylabel('Predicted Activity')
plot.savefig('TestSet.png')
plot.savefig('TestSet.png')
plot.show()


R1="{:.2f}".format(np.corrcoef(yTest, yPredictionTest)[0,1])
R2="{:.2f}".format(np.corrcoef(yTrain, yPredictionTrain)[0,1])

MAE1="{:.2f}".format(mean_absolute_error(yTest, yPredictionTest))
MAE2="{:.2f}".format(mean_absolute_error(yTrain, yPredictionTrain))

RMSE1="{:.2f}".format(mean_squared_error(yTest, yPredictionTest,squared=False))
RMSE2="{:.2f}".format(mean_squared_error(yTrain, yPredictionTrain,squared=False))

R1=pd.Series(R1)
R1=R1.rename('R1')
R2=pd.Series(R2)
R2=R2.rename('R2')
MAE1=pd.Series(MAE1)
MAE1=MAE1.rename('MAE1')
MAE2=pd.Series(MAE2)
MAE2=MAE2.rename('MAE2')
RMSE1=pd.Series(RMSE1)
RMSE1=RMSE1.rename('RMSE1')
RMSE2=pd.Series(RMSE2)
RMSE2=RMSE2.rename('RMSE2')

data = {"R-Test": R1, 
        "R-Train": R2,
        "MAE-Test":MAE1,
        "MAE-Train":MAE2,
        "RMSE-Test":RMSE1,
        "RMSE-Train":RMSE2} 
DF=pd.DataFrame(data)

from docx import Document
from docx.shared import Inches

document = Document()

#p = document.add_paragraph()
#r = p.add_run()
#r.add_text('Evaluation Report of Linear Regression Model')


p = document.add_paragraph()
r = p.add_run()
r.add_picture('TestSet.png')

#p = document.add_paragraph()
#r = p.add_run()
r.add_picture('TrainSet.png')

p = document.add_paragraph()
r = p.add_run()
r.add_text("Correlation Co-efficient Value on Test Set (R-Test): "+"{:.2f}".format(np.corrcoef(yTest, yPredictionTest)[0,1]))

p = document.add_paragraph()
r = p.add_run()
r.add_text("Correlation Co-efficient Value on Train Set (R-Train): "+"{:.2f}".format(np.corrcoef(yTrain, yPredictionTrain)[0,1]))


p = document.add_paragraph()
r = p.add_run()
r.add_text("Mean Absolute Error on Test Set: "+"{:.2f}".format(mean_absolute_error(yTest, yPredictionTest))) 

p = document.add_paragraph()
r = p.add_run()
r.add_text("Mean Absolute Error on Train Set: "+"{:.2f}".format(mean_absolute_error(yTrain, yPredictionTrain)))

p = document.add_paragraph()
r = p.add_run()
r.add_text("Root Mean Squared Error (RMSE) on Test Set: "+"{:.2f}".format(mean_squared_error(yTest, yPredictionTest,squared=False))) 

p = document.add_paragraph()
r = p.add_run()
r.add_text("Root Mean Squared Error (RMSE) on Train Set: "+"{:.2f}".format(mean_squared_error(yTrain, yPredictionTrain,squared=False))) 


document.save('LR-Report.docx')

